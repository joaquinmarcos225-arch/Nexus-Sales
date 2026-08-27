"""Extrae texto de documentos de producto (PDF, DOCX, texto plano, etc.)."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass


class DocumentExtractError(Exception):
    def __init__(self, message: str):
        self.message = message
        super().__init__(message)


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    filename: str
    format: str
    chars: int


_MAX_CHARS = 480_000

_TEXT_EXT = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".tsv",
    ".json",
    ".html",
    ".htm",
    ".xml",
    ".rtf",
    ".log",
}


def _clean_text(raw: str) -> str:
    text = (raw or "").replace("\x00", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    if len(text) > _MAX_CHARS:
        text = text[:_MAX_CHARS].rstrip() + "\n…"
    return text


def _ext(filename: str) -> str:
    name = (filename or "").strip().lower()
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[-1]


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractError(
            "Falta dependencia pypdf. Ejecutá: pip install pypdf"
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise DocumentExtractError(f"No se pudo leer el PDF: {exc}") from exc
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(parts)


def _from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractError(
            "Falta dependencia python-docx. Ejecutá: pip install python-docx"
        ) from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentExtractError(f"No se pudo leer el DOCX: {exc}") from exc
    parts = [p.text for p in doc.paragraphs if (p.text or "").strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_text_bytes(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_document_text(*, filename: str, data: bytes, content_type: str | None = None) -> ExtractedDocument:
    if not data:
        raise DocumentExtractError("El archivo está vacío.")
    if len(data) > 12 * 1024 * 1024:
        raise DocumentExtractError("El archivo supera 12 MB.")

    ext = _ext(filename)
    ctype = (content_type or "").split(";")[0].strip().lower()

    if ext == ".pdf" or ctype == "application/pdf":
        text = _clean_text(_from_pdf(data))
        fmt = "pdf"
    elif ext in {".docx"} or ctype in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        text = _clean_text(_from_docx(data))
        fmt = "docx"
    elif ext == ".doc":
        raise DocumentExtractError(
            "Los .doc antiguos no están soportados. Guardá como .docx o PDF."
        )
    elif ext in _TEXT_EXT or ctype.startswith("text/") or ctype in {
        "application/json",
        "application/xml",
    }:
        text = _clean_text(_from_text_bytes(data))
        fmt = ext.lstrip(".") or "text"
    else:
        # Intento genérico: si es texto legible, úsalo; si no, rechazar con hint.
        probe = _from_text_bytes(data[:4000])
        printable = sum(1 for ch in probe if ch.isprintable() or ch in "\n\r\t")
        if probe and printable / max(len(probe), 1) > 0.85:
            text = _clean_text(_from_text_bytes(data))
            fmt = ext.lstrip(".") or "bin-text"
        else:
            raise DocumentExtractError(
                "Formato no soportado. Usá PDF, DOCX, TXT, MD, CSV, HTML o JSON."
            )

    if len(text) < 40:
        raise DocumentExtractError(
            "No se pudo extraer texto suficiente del archivo (mínimo ~40 caracteres). "
            "Si es un PDF escaneado (imagen), convertí a texto o pegá el contenido."
        )

    return ExtractedDocument(
        text=text,
        filename=(filename or "documento")[:255],
        format=fmt,
        chars=len(text),
    )
